# -*- coding: UTF-8 -*-
import datetime
from gnr.core.gnrdecorator import public_method
from gnr.core.gnrbag import Bag
from docx import Document

class GnrCustomWebPage(object):

    def main(self,root,**kwargs):
        pane =root.contentPane(height='100%', margin='15px', border='1px solid silver', datapath='test')

        fb = pane.formbuilder(cols=2)
        fb.dbselect(value='^.cliente_id',dbtable='fatt.cliente', lbl='Cliente')
        fb.textbox( value='^.titolo',lbl='Titolo')
        fb.simpleTextArea(colspan=2, value='^.testo', width='400px', height='300px', lbl='Testo offerta')
        fb.button('Genera doc offerta', action='FIRE generaDoc;', disabled='^.cliente_id?=!#v')
        pane.dataRpc('.result', self.generaDoc, _fired='^generaDoc', 
                    cliente_id='=.cliente_id', titolo='=.titolo',
                    testo='=.testo')

    @public_method
    def generaDoc(self, cliente_id=None, titolo=None, testo=None):
        cliente_record = self.db.table('fatt.cliente').record(cliente_id).output('bag')
        document = Document()

        document.add_heading(titolo, 0)

        document.add_heading('Milano:{data} alla cortese attenzione di {cliente}'.format(data= self.workdate, cliente = cliente_record['ragione_sociale']), level=1)

        par = document.add_paragraph()
        par.add_run(text=testo)
        filename = 'offerta_{cliente}.docx'.format(cliente=cliente_record['ragione_sociale'].replace(' ','_'))
        sn = self.site.storageNode('page:documenti',filename)
        with sn.local_path() as path:
            document.save(path)
        self.setInClientData(path='gnr.downloadurl',value=sn.url(),fired=True)
                                